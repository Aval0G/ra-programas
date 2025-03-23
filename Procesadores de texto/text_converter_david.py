import os
import tkinter as tk
from tkinter import Tk, Label, Button, filedialog, StringVar, OptionMenu, filedialog, StringVar, ttk
from docx import Document as DocxDocument
from odf.opendocument import OpenDocumentText, load
from odf.text import P
from striprtf.striprtf import rtf_to_text
from bs4 import BeautifulSoup
import PyPDF2
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas
import textract
from PyRTF.Elements import Document as RTFDocument, Section as RTFSection, Text as RTFText
from PyRTF.Renderer import Renderer, Paragraph as RTFParagraph
import pypandoc
import re

SUPPORTED_EXTENSIONS = ['.txt', '.doc', '.docx', '.pdf', '.rtf', '.html', '.odt']

"""READ FUNCTIONS"""
def read_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def read_docx(file_path):
    doc = DocxDocument(file_path)
    return '\n'.join(p.text for p in doc.paragraphs)

def read_pdf(file_path):
    text = ""
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text()
    return text

def read_rtf(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return rtf_to_text(f.read())

def read_html(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f, 'html.parser')
        return soup.get_text()

def read_odt(file_path):
    doc = load(file_path)
    text = ""
    for paragraph in doc.getElementsByType(P):
        for node in paragraph.childNodes:
            if node.nodeType == node.TEXT_NODE:
                text += node.data
        text += '\n'
    return text

def read_doc(file_path):
    try:
        text = textract.process(file_path).decode('utf-8')
        return text
    except Exception as e:
        return f"Error reading DOC file: {e}"

"""WRITE FUNCTIONS"""
def write_txt(text, file_path):
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(text)

def write_docx(text, file_path):
    doc = DocxDocument()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    doc.save(file_path)

def write_odt(text, file_path):
    doc = OpenDocumentText()
    doc.text.addElement(P(text=text))
    doc.save(file_path)

def write_pdf(text, file_path):
    pdf = canvas.Canvas(file_path, pagesize=LETTER)
    width, height = LETTER
    y = height - 40
    for line in text.split('\n'):
        pdf.drawString(40, y, line)
        y -= 15
        if y < 40:
            pdf.showPage()
            y = height - 40
    pdf.save()

def write_rtf(text, file_path):
    def escape_rtf(txt):
        txt = txt.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
        txt = re.sub(r'([\u0080-\uffff])', lambda m: r'\u{}?'.format(ord(m.group(1))), text)
        return txt

    doc = RTFDocument()
    section = RTFSection()
    doc.Sections.append(section)
    
    paragraph = RTFParagraph()
    paragraph.append(RTFText(escape_rtf(text)))
    section.append(paragraph)

    renderer = Renderer()
    
    with open(file_path, 'w', encoding='utf-8') as fout:
        renderer.Write(doc, fout)

def write_html(text, file_path):
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write('<html><body>')
        for paragraph in text.split('\n\n'):
            f.write('<p>')
            f.write(paragraph.replace('\n', '<br>'))
            f.write('</p>')
        f.write('</body></html>')

READERS = {
    '.txt': read_txt,
    '.docx': read_docx,
    '.pdf': read_pdf,
    '.rtf': read_rtf,
    '.html': read_html,
    '.odt': read_odt,
    '.doc': read_doc
}

WRITERS = {
    '.txt': write_txt,
    '.docx': write_docx,
    '.pdf': write_pdf,
    '.rtf': write_rtf,
    '.html': write_html,
    '.odt': write_odt
}

def convert_file(input_file, output_format, output_folder):
    ext = os.path.splitext(input_file)[1].lower()
    reader = READERS.get(ext)
    
    if not reader:
        print(f"No reader implemented for '{ext}' files.")
        return 1
    
    # Manage error when reading the file
    text = ""
    try:
        text = reader(input_file)
    except Exception as e:
        print(f"Error reading file: {e}")
        return 1

    output_file = os.path.join(output_folder, f"{os.path.splitext(os.path.basename(input_file))[0]}_converted{output_format}")
    
    writer = WRITERS.get(output_format)
    
    if not writer:
        print(f"No writer implemented for '{output_format}' files.")
        return 1
    
    # Manage error when writing the file
    try:
        writer(text, output_file)
    except Exception as e:
        print(f"Error writing file: {e}")
        return 1
    
    print(f"Conversion complete! File saved to: {output_file}")

    return 0

def select_input_file(input_file_var):
    file_path = filedialog.askopenfilename(filetypes=[("All Files", "*.*")])
    if file_path:
        input_file_var.set(file_path)

def select_output_folder(output_folder_var):
    folder_path = filedialog.askdirectory()
    if folder_path:
        output_folder_var.set(folder_path)

def start_conversion(input_file_var, output_folder_var, output_format_var, status_label):
    input_file = input_file_var.get()
    output_folder = output_folder_var.get()
    output_format = output_format_var.get()

    if not input_file or not output_folder:
        status_label.config(text="ERROR: Select both input file and output folder!", foreground="red")
        return

    input_file_name = os.path.basename(input_file)
    error = convert_file(input_file, output_format, output_folder)  # Call the conversion function
    
    if error:
        status_label.config(text="ERROR: Conversion failed!", foreground="red")
    else:
        status_label.config(text=f"SUCCESS: Converted {input_file_name} to {output_format}!", foreground="green")
        # Open the output folder in the file explorer
        os.startfile(output_folder)

def main():
    # Create the main application window
    app = tk.Tk()
    app.title("Text Converter")

    # Set window size
    width, height = 700, 400

    # Get screen width and height
    screen_width = app.winfo_screenwidth()
    screen_height = app.winfo_screenheight()

    # Calculate x and y coordinates for the window to be centered
    x = (screen_width // 2) - (width // 2)
    y = (screen_height // 2) - (height // 2)

    # Set geometry with calculated position
    app.geometry(f"{width}x{height}+{x}+{y}")
    app.configure(bg="#f0f0f0")

    # Variables
    input_file_var = StringVar()
    output_folder_var = StringVar()
    output_format_var = StringVar(value=".pdf")

    # Main frame using pack instead of grid
    main_frame = ttk.Frame(app, padding=20)
    main_frame.pack(expand=True)

    # File Selection
    ttk.Label(main_frame, text="Select Input File:").pack(anchor="w", pady=5)
    file_entry = ttk.Entry(main_frame, textvariable=input_file_var, width=40, font=("Arial", 12), state='readonly')
    file_entry.pack(pady=5)
    ttk.Button(main_frame, text="Browse", command=lambda: select_input_file(input_file_var)).pack(pady=5)

    # Output Folder Selection
    ttk.Label(main_frame, text="Select Output Folder:").pack(anchor="w", pady=5)
    folder_entry = ttk.Entry(main_frame, textvariable=output_folder_var, width=40, font=("Arial", 12), state='readonly')
    folder_entry.pack(pady=5)
    ttk.Button(main_frame, text="Browse", command=lambda: select_output_folder(output_folder_var)).pack(pady=5)

    # Output Format Selection
    ttk.Label(main_frame, text="Select Output Format:").pack(anchor="w", pady=5)
    format_dropdown = ttk.Combobox(main_frame, textvariable=output_format_var, values=[".pdf", ".docx", ".odt", ".txt", ".rtf", ".html"], font=("Arial", 12), state='readonly')
    format_dropdown.pack(pady=5)
    format_dropdown.current(0)

    # Convert Button
    convert_button = ttk.Button(main_frame, text="Convert", command=lambda: start_conversion(input_file_var, output_folder_var, output_format_var, status_label))
    convert_button.pack(pady=20)

    # Status Label
    status_label = ttk.Label(main_frame, text="", font=("Arial", 12, "bold"))
    status_label.pack(pady=5)

    app.mainloop()

if __name__ == "__main__":
    main()
